"""
Build SARA_MESSAGE_SCRIPTS.docx from WHATSAPP_SCRIPTS.md (Hebrew, RTL, tables).
Run from repo root: python scripts/generate_sara_message_scripts.py
"""

from __future__ import annotations

import re
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def add_paragraph_bidi(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for el in list(p_pr):
        if el.tag == qn("w:bidi"):
            p_pr.remove(el)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_runs_with_variables(paragraph, text: str) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)
    add_paragraph_bidi(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    parts = re.split(r"(\[[^\]]+\])", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(11)
        run.font.name = "Arial"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Arial")
        rfonts.set(qn("w:hAnsi"), "Arial")
        rfonts.set(qn("w:cs"), "Arial")
        if part.startswith("[") and part.endswith("]"):
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)


def parse_rows(md: str) -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []
    chunks = re.split(r"\n(?=### )", md)
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk.startswith("###"):
            continue
        first_line, _, rest = chunk.partition("\n")
        title = first_line.replace("###", "").strip()
        if title.startswith("הערה:") and "ארכיון" in title:
            m = re.search(
                r"הטקסט שהיה מתוכנן \(ארכיון\):\s*\n\n```\n(.*?)```",
                rest,
                re.DOTALL,
            )
            if m:
                rows.append((title, "קוד מושבת — לא נשלח לעובד", m.group(1).strip()))
            continue
        trig_m = re.search(r"\*\*TRIGGER:\*\*\s*(.+)", rest)
        alt_m = re.search(r"\*\*נשלח אל:\*\*\s*(.+)", rest)
        if trig_m:
            trig = trig_m.group(1).strip()
        elif alt_m:
            trig = alt_m.group(1).strip()
        else:
            trig = ""

        for m in re.finditer(
            r"\*\*MESSAGE[^*]*\*\*[^\n]*\n\n```\n(.*?)```",
            rest,
            re.DOTALL,
        ):
            msg = m.group(1).strip()
            rows.append((title, trig, msg))
    return rows


def section_key(title: str) -> str:
    if title.startswith("SMS") or title.startswith("SMS "):
        return "sms"
    if "בדיקת חיבור" in title or ("הערה:" in title and "ארכיון" in title):
        return "extra"
    return "wa"


def set_table_rtl(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    bidi_visual = OxmlElement("w:bidiVisual")
    bidi_visual.set(qn("w:val"), "1")
    tbl_pr.append(bidi_visual)


def add_table_header(table) -> None:
    hdr = table.rows[0].cells
    headers = ["שם זרימה / כותרת", "תנאי טריגר (מפושט)", "הודעה מקורית", "הודעה מעודכנת"]
    widths = [Cm(3.0), Cm(3.5), Cm(5.8), Cm(5.8)]
    for i, ht in enumerate(headers):
        p = hdr[i].paragraphs[0]
        add_paragraph_bidi(p)
        r = p.add_run(ht)
        r.bold = True
        r.font.size = Pt(11)
        hdr[i].width = widths[i]


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "WHATSAPP_SCRIPTS.md"
    out_path = root / "SARA_MESSAGE_SCRIPTS.docx"
    if not md_path.exists():
        print("Missing WHATSAPP_SCRIPTS.md", file=sys.stderr)
        return 1

    md = md_path.read_text(encoding="utf-8")
    rows = parse_rows(md)
    if not rows:
        print("No rows parsed", file=sys.stderr)
        return 1

    grouped: dict[str, list[tuple[str, str, str]]] = defaultdict(list)
    for row in rows:
        grouped[section_key(row[0])].append(row)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("סקריפט הודעות — עריכה לסרה", level=0)
    add_paragraph_bidi(title)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    intro = doc.add_paragraph()
    add_runs_with_variables(
        intro,
        "מסמך זה מרכז את כל הודעות הוואטסאפ לדיירים והודעות ה־SMS לצוות. "
        "בעמודה «הודעה מעודכנת» ניתן לכתוב את הניסוח החדש. "
        "טקסטים בסוגריים מרובעים [כך] הם משתנים מהמערכת — הודגשו בכחול ובולד.",
    )
    doc.add_paragraph()

    section_order = [
        ("wa", "חלק א׳ — הודעות וואטסאפ לדיירים"),
        ("extra", "חלק ב׳ — הודעות נוספות (בדיקות / ארכיון)"),
        ("sms", "חלק ג׳ — הודעות SMS למנהל ולעובדים"),
    ]

    for key, label in section_order:
        if key not in grouped:
            continue
        h = doc.add_heading(label, level=1)
        add_paragraph_bidi(h)
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        table = doc.add_table(rows=1, cols=4)
        set_table_rtl(table)
        add_table_header(table)

        for flow_title, trigger, message in grouped[key]:
            cells = table.add_row().cells
            add_runs_with_variables(cells[0].paragraphs[0], flow_title)
            add_runs_with_variables(cells[1].paragraphs[0], trigger or "—")
            add_runs_with_variables(cells[2].paragraphs[0], message)
            p3 = cells[3].paragraphs[0]
            add_paragraph_bidi(p3)
            p3.add_run("\u00a0").font.size = Pt(11)

        doc.add_paragraph()

    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[0]
        fp.text = ""
        add_paragraph_bidi(fp)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run("לאחר עריכה — שלח את המסמך חזרה לצוות הטכני")
        fr.italic = True
        fr.font.size = Pt(10)

    doc.save(out_path)
    print(f"Wrote {out_path} ({len(rows)} message rows in {len(grouped)} sections)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
